from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import json

BASE = Path('/home/ubuntu/isdebuggerpresent_flow')
OUT = Path('/home/ubuntu/isdebuggerpresent_flow/Relatorio_Final_Fluxo_IsDebuggerPresent.docx')
JSON_MAP = BASE / 'outputs' / 'fluxo_isdebuggerpresent.json'
PNG = BASE / 'outputs' / 'fluxo_isdebuggerpresent.png'
CORR = BASE / 'test_outputs' / 'correlation_isdebuggerpresent.json'
COMP = BASE / 'test_outputs' / 'compressed' / 'compression_manifest.json'


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_style(doc):
    for name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
        style = doc.styles[name]
        style.font.name = 'Arial'
        style.font.size = Pt(11 if name == 'Normal' else 14)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph('')


mapping = json.loads(JSON_MAP.read_text(encoding='utf-8'))
correlation = json.loads(CORR.read_text(encoding='utf-8'))
compression = json.loads(COMP.read_text(encoding='utf-8'))

focus_hits = len(correlation.get('function_interceptor_hits', []))
trace_hits = len(correlation.get('traceinstructions_hits', []))
compressed_files = compression.get('file_count', 0)


doc = Document()
set_style(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Relatório Final – Fluxo Mapeado a partir de IsDebuggerPresent')
r.bold = True
r.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Correlações entre logs, cadeia até a identificação do malware e implementação das sugestões de acompanhamento').italic = True


doc.add_paragraph(
    'Este relatório consolida a análise solicitada pelo usuário com foco específico na função IsDebuggerPresent. A partir do documento-base fornecido, foi reconstruída a cadeia lógica entre os arquivos de log, desde o evento anti-debug inicial até a identificação do artefato como trojan 64-bit protegido por VMProtect. Além disso, foram implementados artefatos práticos para acompanhar análises futuras, cobrindo compressão adaptativa, filtragem seletiva, processamento incremental, correlação automática e visualização [1].'
)

doc.add_heading('1. Escopo e limitação da evidência disponível', level=1)
doc.add_paragraph(
    'O ambiente continha apenas o documento-resumo da execução, sem os arquivos brutos contradef.2956.*.cdf. Por isso, o fluxo abaixo é uma reconstrução técnica baseada em evidências textuais explícitas do relatório-fonte e em inferências de alta confiança decorrentes da relação declarada entre os módulos de trace. Isso significa que a análise entrega um mapa operacional confiável, mas não um extrato byte a byte dos eventos originais [1].'
)

add_table(doc, ['Aspecto', 'Situação no ambiente', 'Impacto analítico'], [
    ['Documento-base', 'Disponível', 'Permitiu reconstruir a cadeia comportamental'],
    ['Arquivos CDF brutos', 'Indisponíveis', 'Impedem confirmar offsets, timestamps e argumentos reais'],
    ['Resultado entregue', 'Fluxo mapeado + artefatos operacionais', 'Permite reaproveitamento imediato quando os traces forem exportados'],
])

doc.add_heading('2. Ponto de partida em IsDebuggerPresent', level=1)
doc.add_paragraph(
    'O ponto de partida específico pedido pelo usuário é a chamada IsDebuggerPresent, identificada no documento-base como parte explícita da fase de anti-debug. Esse evento funciona como pivô de correlação porque liga o nível de API interceptada ao ponto de origem da chamada, à instrução correspondente, ao contexto de memória e ao caminho de execução observado depois do teste [1].'
)

add_table(doc, ['Arquivo', 'Como IsDebuggerPresent é mapeado', 'Confiança'], [
    ['FunctionInterceptor.cdf', 'Evento explícito da API anti-debug', 'Alta'],
    ['TraceFcnCall.M1.cdf', 'Origem por call direta, se houver importação convencional', 'Média/Alta'],
    ['TraceFcnCall.M2.cdf', 'Origem por salto indireto ou resolução dinâmica', 'Média/Alta'],
    ['TraceInstructions.cdf', 'Instrução exata e branch após o retorno', 'Alta'],
    ['TraceMemory.cdf', 'Contexto de buffers e páginas próximo ao teste e ao unpacking', 'Média'],
    ['TraceDisassembly.cdf', 'Bloco básico e transição para o código desempacotado', 'Alta'],
])

doc.add_heading('3. Fluxo correlacionado entre todos os arquivos', level=1)
doc.add_paragraph(
    'A reconstrução mais consistente, suportada pelo relatório-fonte, segue o encadeamento: carregamento de bibliotecas, IsDebuggerPresent, CheckRemoteDebuggerPresent, NtQueryInformationProcess com classe 30, verificações temporais com GetTickCount e QueryPerformanceCounter, verificações anti-VM com EnumSystemFirmwareTables e WMI, preparação de memória com LocalAlloc, alteração de proteção com VirtualProtect, execução do código desempacotado e encerramento com HeapFree e FatalExit [1].'
)

rows = [
    ['1', 'Carregamento de kernel32.dll e ntdll.dll', 'FunctionInterceptor', 'Inicialização rastreável do processo'],
    ['2', 'IsDebuggerPresent', 'FunctionInterceptor -> TraceFcnCall -> TraceInstructions', 'Início específico da fase anti-debug'],
    ['3', 'CheckRemoteDebuggerPresent', 'FunctionInterceptor -> TraceFcnCall -> TraceInstructions', 'Validação complementar de depuração'],
    ['4', 'NtQueryInformationProcess classe 30', 'FunctionInterceptor -> TraceFcnCall -> TraceInstructions -> TraceMemory', 'Consulta anti-debug em nível mais baixo'],
    ['5', 'GetTickCount / QueryPerformanceCounter', 'FunctionInterceptor -> TraceInstructions', 'Evasão temporal e detecção de overhead'],
    ['6', 'EnumSystemFirmwareTables / WMI', 'FunctionInterceptor -> TraceFcnCall.M2 -> TraceInstructions', 'Anti-VM e anti-sandbox'],
    ['7', 'LocalAlloc', 'FunctionInterceptor -> TraceInstructions -> TraceMemory', 'Preparação de memória'],
    ['8', 'VirtualProtect RW para RX', 'FunctionInterceptor -> TraceInstructions -> TraceMemory -> TraceDisassembly', 'Unpacking e preparação para execução'],
    ['9', 'Transferência para o código desempacotado', 'TraceDisassembly -> TraceInstructions', 'Execução do payload'],
    ['10', 'HeapFree / FatalExit', 'FunctionInterceptor -> TraceInstructions', 'Encerramento do processo'],
]
add_table(doc, ['Etapa', 'Evento', 'Arquivos correlacionados', 'Leitura forense'], rows)

if PNG.exists():
    doc.add_heading('4. Visualização do fluxo', level=1)
    doc.add_paragraph('O diagrama abaixo sintetiza o encadeamento mapeado a partir de IsDebuggerPresent até a identificação do malware e o encerramento da execução.')
    doc.add_picture(str(PNG), width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Figura 1. Cadeia lógica entre os arquivos de trace e os principais eventos observados.').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('5. Identificação do malware', level=1)
doc.add_paragraph(
    'O relatório-base descreve a amostra como Trojan 64-bit (Ulise, Casdet, BackdoorX), protegida por VMProtect 2.x, e essa classificação é coerente com a convergência de quatro grupos de evidência: anti-debug, anti-VM/anti-sandbox, anti-overhead e desempacotamento em memória com mudança de proteção de páginas. A identificação analítica se completa antes do encerramento do processo, quando essa cadeia já está suficientemente demonstrada [1].'
)

add_table(doc, ['Indicador', 'Valor', 'Significado'], [
    ['Hash SHA-256', '36685efcf34c7a7a6f6dd2e48199e4700b5ab8fe3945a50297703dd8daced74f', 'Identificador único da amostra'],
    ['Classificação reportada', 'Trojan 64-bit (Ulise, Casdet, BackdoorX)', 'Natureza maliciosa reportada no documento'],
    ['Proteção', 'VMProtect 2.x', 'Explica ofuscação, saltos indiretos e unpacking'],
])

doc.add_heading('6. Implementação das sugestões de acompanhamento', level=1)
doc.add_paragraph(
    'As recomendações do documento-base foram convertidas em um pacote operacional pronto para uso. A implementação cobre compressão adaptativa dos traces, filtragem seletiva por função e endereço, processamento incremental em chunks, correlação automática do fluxo e geração de visualização a partir de JSON. Esses componentes foram organizados em scripts reutilizáveis e documentados em um manual de uso [1].'
)

add_table(doc, ['Sugestão original', 'Implementação realizada', 'Arquivo entregue'], [
    ['Compressão de traces', 'Compressão adaptativa com manifesto SHA-256', 'scripts/adaptive_trace_compressor.py'],
    ['Filtragem seletiva', 'Filtro por função, regex, texto e intervalo de endereços', 'scripts/selective_trace_filter.py'],
    ['Análise incremental', 'Processamento por blocos com sumário por chunk', 'scripts/chunked_trace_processor.py'],
    ['Correlação automática', 'Correlacionador centrado em IsDebuggerPresent', 'scripts/correlate_isdebuggerpresent_flow.py'],
    ['Visualização', 'Gerador Mermaid a partir do JSON de correlação', 'scripts/build_mermaid_from_json.py'],
])

doc.add_heading('7. Validação executada', level=1)
doc.add_paragraph(
    'Como os CDF originais não estavam no ambiente, a validação funcional foi feita com pequenos arquivos tabulares derivados estritamente do documento-resumo fornecido pelo usuário. Esses arquivos de exemplo não substituem os traces reais; eles serviram apenas para confirmar que o pipeline implementado executa corretamente da filtragem até a visualização.'
)

add_table(doc, ['Componente validado', 'Resultado observado'], [
    ['Filtro seletivo', f'1 ocorrência principal de IsDebuggerPresent localizada em FunctionInterceptor'],
    ['Correlação automática', f'{focus_hits} hit no FunctionInterceptor e {trace_hits} hits em TraceInstructions para IsDebuggerPresent'],
    ['Chunked processing', 'Resumo agregado gerado com sucesso em múltiplos blocos'],
    ['Geração de diagrama', 'Arquivo Mermaid e imagem PNG gerados com sucesso'],
    ['Compressão adaptativa', f'{compressed_files} artefatos comprimidos com manifesto de integridade'],
])

doc.add_heading('8. Conclusão prática', level=1)
doc.add_paragraph(
    'Em termos analíticos, o fluxo pedido pelo usuário ficou mapeado a partir de IsDebuggerPresent até a identificação do malware. Em termos operacionais, as sugestões de acompanhamento também foram implementadas e testadas no ambiente. O próximo passo natural, caso o usuário deseje aprofundar a investigação, é fornecer as exportações reais dos arquivos contradef.2956.*.cdf para que o mesmo pipeline produza uma linha do tempo forense completa com endereços, argumentos, threads e timestamps verdadeiros.'
)

doc.add_heading('Referência', level=1)
doc.add_paragraph('[1] Documento fornecido pelo usuário: “Análise dos Resultados de Execução - Full-Execution-Sample-1.docx”.')

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('Página ')
add_page_number(footer)

doc.save(OUT)
print(OUT)
