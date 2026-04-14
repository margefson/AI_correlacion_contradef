#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


def load(path: str | Path):
    return json.loads(Path(path).read_text(encoding='utf-8'))


def format_bytes(num: int) -> str:
    units = ['B', 'KB', 'MB', 'GB', 'TB']
    value = float(num)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            return f'{value:.2f} {unit}'
        value /= 1024
    return f'{num} B'


def add_table(document: Document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row.get(h, ''))
    return table


def main() -> None:
    parser = argparse.ArgumentParser(description='Gerar relatório DOCX da análise dos CDFs reais.')
    parser.add_argument('--repo-root', required=True)
    parser.add_argument('--output-docx', required=True)
    args = parser.parse_args()

    root = Path(args.repo_root)
    dataset_manifest = load(root / 'data/manifests/dataset_manifest.json')
    compression_manifest = load(root / 'data/manifests/compression_manifest.json')
    flow = load(root / 'data/correlation/isdebuggerpresent_flow_real.json')
    ti = load(root / 'data/derived/traceinstructions_chunk_summary.json')
    tm = load(root / 'data/derived/tracememory_chunk_summary.json')

    doc = Document()
    doc.core_properties.author = 'Manus AI'
    doc.core_properties.title = 'Relatório Final - CDFs Reais e Fluxo IsDebuggerPresent'

    doc.add_heading('Relatório final dos CDFs reais: fluxo a partir de IsDebuggerPresent', level=0)
    doc.add_paragraph(
        'Este documento consolida a execução do pipeline completo sobre os CDFs reais do pacote Full-Execution-Sample-1. '
        'O objetivo foi reconstruir o fluxo centrado em IsDebuggerPresent, implementar as sugestões de acompanhamento '
        'sobre dados reais e preservar artefatos reutilizáveis no repositório do projeto.'
    )
    doc.add_paragraph(
        'O fluxo agregado observado parte de LoadLibraryA e IsDebuggerPresent, evolui por verificações anti-debug, consultas '
        'de processo, marcação temporal e manipulação de memória, e termina no nó FatalExit no encadeamento consolidado.'
    )

    doc.add_heading('Dataset utilizado', level=1)
    dataset_rows = []
    for item in dataset_manifest:
        dataset_rows.append({
            'Arquivo': item['file'],
            'Tamanho': format_bytes(item['size_bytes']),
            'SHA-256': item['sha256'][:16] + '...',
        })
    add_table(doc, ['Arquivo', 'Tamanho', 'SHA-256'], dataset_rows)

    doc.add_heading('Compressão adaptativa executada', level=1)
    doc.add_paragraph(
        'Como acompanhamento operacional, os CDFs reais completos foram comprimidos localmente com nível adaptativo, '
        'e o resultado foi registrado em manifesto versionado. Essa etapa viabiliza preservação e redistribuição controlada '
        'dos traces sem reprocessamento inicial.'
    )
    compression_rows = []
    for item in compression_manifest['artifacts']:
        ratio = 0 if item['original_size'] == 0 else (1 - (item['compressed_size'] / item['original_size'])) * 100
        compression_rows.append({
            'Arquivo': Path(item['source']).name,
            'Original': format_bytes(item['original_size']),
            'Comprimido': format_bytes(item['compressed_size']),
            'Redução (%)': f'{ratio:.2f}',
            'Nível': item['compression_level'],
        })
    add_table(doc, ['Arquivo', 'Original', 'Comprimido', 'Redução (%)', 'Nível'], compression_rows)

    doc.add_heading('Fluxo correlacionado', level=1)
    doc.add_paragraph(
        'A tabela seguinte resume as contagens por função em cada família de log: FunctionInterceptor (FI), TraceFcnCall.M1 (M1), '
        'TraceFcnCall.M2 (M2), TraceInstructions (TI) e TraceMemory (TM).'
    )
    flow_rows = []
    for node in flow['nodes']:
        flow_rows.append({
            'Função': node['label'],
            'FI': node['counts']['function_interceptor'],
            'M1': node['counts']['tracefcncall_m1'],
            'M2': node['counts']['tracefcncall_m2'],
            'TI': node['counts']['traceinstructions'],
            'TM': node['counts']['tracememory'],
        })
    add_table(doc, ['Função', 'FI', 'M1', 'M2', 'TI', 'TM'], flow_rows)

    figure_path = root / 'data/figures/isdebuggerpresent_flow_real.png'
    if figure_path.exists():
        doc.add_heading('Diagrama do fluxo', level=1)
        doc.add_paragraph('O diagrama abaixo resume o encadeamento reconstruído e a presença relativa de cada função nos diferentes arquivos de trace.')
        doc.add_picture(str(figure_path), width=Inches(4.5))

    doc.add_heading('Processamento incremental por chunks', level=1)
    doc.add_paragraph(
        f'O TraceInstructions foi analisado em streaming sobre {ti["total_lines"]:,} linhas, enquanto o TraceMemory foi processado sobre {tm["total_lines"]:,} linhas. '
        'Esse mecanismo reduz custo de memória e torna possível reexecutar filtros seletivos e correlação de forma incremental.'
    )
    ti_rows = []
    for chunk in ti['chunks'][:8]:
        if chunk['term_counts']:
            ti_rows.append({
                'Chunk': chunk['chunk_index'],
                'Linhas': f"{chunk['start_line']}-{chunk['end_line']}",
                'Ocorrências': ', '.join(f"{k}:{v}" for k, v in chunk['term_counts'].items()),
            })
    if ti_rows:
        add_table(doc, ['Chunk', 'Linhas', 'Ocorrências'], ti_rows)

    doc.add_heading('Chamadas-chave extraídas', level=1)
    doc.add_paragraph(
        'O correlacionador preservou chamadas e tailcalls relevantes no TraceInstructions. As primeiras arestas abaixo mostram a transição entre símbolos internos e APIs-alvo associadas ao fluxo monitorado.'
    )
    key_rows = []
    for item in flow.get('key_call_edges', [])[:12]:
        key_rows.append({
            'Linha': item['line_number'],
            'Tipo': item['calltype'],
            'Origem': item['src_symbol'][:72],
            'Destino': item['dst_symbol'][:72],
            'Thread': item['thread'],
        })
    if key_rows:
        add_table(doc, ['Linha', 'Tipo', 'Origem', 'Destino', 'Thread'], key_rows)

    doc.add_heading('Artefatos implementados', level=1)
    doc.add_paragraph(
        'Foram implementados no repositório scripts para processamento dos CDFs reais, compressão adaptativa, reconstrução do diagrama Mermaid, geração de relatório em Markdown, geração do presente DOCX e versionamento dos artefatos derivados. Além disso, os manifests, os recortes filtrados, os sumários por chunk e a correlação estruturada ficaram registrados em caminhos previsíveis dentro do projeto.'
    )

    output_path = Path(args.output_docx)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(output_path)


if __name__ == '__main__':
    main()
