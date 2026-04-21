#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from generate_generic_report import build_report


def apply_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_markdown_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def render_markdown_table(document: Document, rows: list[str]) -> None:
    if len(rows) < 2:
        return
    header = parse_markdown_row(rows[0])
    body_rows = [parse_markdown_row(row) for row in rows[2:]]
    table = document.add_table(rows=1 + len(body_rows), cols=len(header))
    table.style = "Table Grid"
    for col_idx, text in enumerate(header):
        table.cell(0, col_idx).text = text
    for row_idx, row in enumerate(body_rows, start=1):
        for col_idx in range(len(header)):
            table.cell(row_idx, col_idx).text = row[col_idx] if col_idx < len(row) else ""


def write_docx(job_dir: Path, output_docx: Path) -> None:
    report_text = build_report(job_dir)
    document = Document()
    apply_default_font(document)
    title = document.add_heading('Relatório de Análise Genérica de CDF', level=0)
    title.alignment = 0

    lines = report_text.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph('')
            index += 1
            continue
        if line.startswith('# '):
            document.add_heading(line[2:].strip(), level=1)
            index += 1
            continue
        if line.startswith('## '):
            document.add_heading(line[3:].strip(), level=2)
            index += 1
            continue
        if line.startswith('### '):
            document.add_heading(line[4:].strip(), level=3)
            index += 1
            continue
        if is_markdown_table_line(line):
            table_lines: list[str] = []
            while index < len(lines) and is_markdown_table_line(lines[index].rstrip()):
                table_lines.append(lines[index].rstrip())
                index += 1
            render_markdown_table(document, table_lines)
            continue
        document.add_paragraph(line)
        index += 1

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='Gerar relatório DOCX a partir de um job da análise genérica de CDF.')
    parser.add_argument('--job-dir', required=True, help='Diretório do job concluído.')
    parser.add_argument('--output-docx', required=True, help='Arquivo DOCX de saída.')
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    job_dir = Path(args.job_dir).expanduser().resolve()
    output_docx = Path(args.output_docx).expanduser().resolve()
    write_docx(job_dir, output_docx)
    print(output_docx)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
